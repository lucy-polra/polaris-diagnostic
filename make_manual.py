from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x00, 0x4B, 0x87)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0xF2, 0xF2, 0xF2)

def bg(cell, hex_color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pr.append(shd)

def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.color.rgb = BLUE
    r.font.bold = True

def tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"💡 {text}")
    r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)
    r.font.italic = True

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

def table_2col(doc, rows, header=True):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        c0.text = a
        c1.text = b
        if i == 0 and header:
            for c in [c0, c1]:
                c.paragraphs[0].runs[0].font.bold = True
                bg(c, "004B87")
                c.paragraphs[0].runs[0].font.color.rgb = WHITE
    doc.add_paragraph()

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3)
sec.right_margin = Cm(2.5)

# ── 표지 ──────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
for text, size, bold in [
    ("폴라리스클라우드", 14, True),
    ("변화관리 진단 에이전트", 28, True),
    ("시스템 사용 매뉴얼", 18, False),
    ("v1.0  |  2026년 6월", 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = BLUE
doc.add_page_break()

# ── 목차 ──────────────────────────────────────────────────
h(doc, "목차")
toc_items = [
    ("1", "시스템 개요"),
    ("2", "시스템 구성"),
    ("3", "접속 방법"),
    ("4", "진단 실행 방법"),
    ("4-1", "1차 사전 진단 설문 (CSV)"),
    ("4-2", "2차 조직 자료 분석 (Excel)"),
    ("4-3", "3차 인터뷰 진단 (JSON)"),
    ("4-4", "4차 행동 진단 (Excel)"),
    ("5", "진단 결과 확인"),
    ("6", "Word 리포트 다운로드"),
    ("7", "코드 수정 방법 (VS Code)"),
    ("8", "GitHub 업데이트 방법"),
    ("9", "자주 묻는 질문 (FAQ)"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5) if "." in num else Cm(0)
    p.add_run(f"{num}.  {title}").font.color.rgb = BLUE
doc.add_page_break()

# ── 1. 시스템 개요 ─────────────────────────────────────────
h(doc, "1. 시스템 개요")
doc.add_paragraph(
    "폴라리스클라우드 변화관리 진단 에이전트는 변화관리 프로젝트 착수 시 "
    "4가지 진단 자료를 입력받아 조직의 변화준비도를 자동으로 분석하고 "
    "Word 리포트를 생성하는 전용 웹 시스템입니다."
)
doc.add_paragraph()
table_2col(doc, [
    ("항목", "내용"),
    ("접속 URL", "https://polaris-diagnostic.streamlit.app"),
    ("접속 비밀번호", "polariscloud2026"),
    ("사용 대상", "폴라리스클라우드 컨설턴트"),
    ("주요 기능", "4가지 진단 자료 분석 + 자동 리포트 생성"),
    ("비용", "무료 (Streamlit Cloud 무료 플랜)"),
])

# ── 2. 시스템 구성 ─────────────────────────────────────────
h(doc, "2. 시스템 구성")
table_2col(doc, [
    ("구성 요소", "역할 및 기술"),
    ("웹 인터페이스", "Streamlit (Python 기반 웹앱)"),
    ("배포 서버", "Streamlit Cloud (무료 호스팅)"),
    ("코드 저장소", "GitHub (lucy-polra/polaris-diagnostic)"),
    ("점수 산출", "규칙 기반 엔진 (AI 미사용, 완전 무료)"),
    ("리포트 생성", "python-docx (Word 자동 생성)"),
    ("데이터 저장", "SharePoint 연동 예정 (M365 E3)"),
])
doc.add_paragraph()
h(doc, "파일 구조", 2)
code(doc, "polaris-diagnostic/")
code(doc, "├── app.py                 ← 메인 (로그인 + 홈)")
code(doc, "├── config.py              ← 환경변수 설정")
code(doc, "├── core/")
code(doc, "│   ├── scoring.py         ← 점수 산출 + 해석 텍스트 + 권고사항")
code(doc, "│   ├── report.py          ← Word 리포트 자동 생성")
code(doc, "│   └── samples.py         ← 샘플 파일 생성")
code(doc, "└── pages/")
code(doc, "    ├── 1_진단실행.py      ← 파일 업로드 + 분석 + 결과")
code(doc, "    └── 2_이력조회.py      ← 진단 이력 조회")
doc.add_page_break()

# ── 3. 접속 방법 ───────────────────────────────────────────
h(doc, "3. 접속 방법")
steps = [
    ("브라우저 접속", "https://polaris-diagnostic.streamlit.app 접속"),
    ("비밀번호 입력", "polariscloud2026 입력 후 로그인 클릭"),
    ("메뉴 선택", "왼쪽 사이드바에서 진단실행 또는 이력조회 선택"),
]
for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"Step {i}. {title}  ")
    r.font.bold = True
    r.font.color.rgb = BLUE
    p.add_run(desc)
tip(doc, "비밀번호는 .env 파일의 APP_PASSWORD 값을 변경하면 수정 가능합니다.")
doc.add_page_break()

# ── 4. 진단 실행 방법 ──────────────────────────────────────
h(doc, "4. 진단 실행 방법")
doc.add_paragraph("왼쪽 사이드바에서 진단실행을 선택한 후 아래 순서대로 진행합니다.")
doc.add_paragraph()

# 4-1
h(doc, "4-1. 1차 사전 진단 설문 (CSV 파일)", 2)
doc.add_paragraph("사람들이 어떻게 생각하는지 파악하는 설문 결과 파일입니다.")
doc.add_paragraph()
table_2col(doc, [
    ("컬럼명", "설명"),
    ("respondent_id", "응답자 ID (예: EMP001)"),
    ("department", "부서명"),
    ("role", "직급 (임원/관리자/실무자)"),
    ("q1_change_necessity", "변화 필요성 공감도 (1~5점)"),
    ("q2_collaboration", "현재 협업 수준 (1~5점)"),
    ("q3_digital_usage", "디지털 활용도 (1~5점)"),
    ("q4_change_fatigue", "변화 피로도 (1~5점, 역산 적용)"),
    ("q5_change_participation", "변화 참여 의지 (1~5점)"),
    ("q6_job_satisfaction", "현재 업무 만족도 (1~5점)"),
    ("q7_communication", "변화 커뮤니케이션 수준 (1~5점)"),
    ("q8_leader_support", "리더의 변화 지원 수준 (1~5점)"),
    ("q9_education_system", "교육 및 지원 체계 수준 (1~5점)"),
    ("q10_change_structure", "변화 추진 체계 인식 (1~5점)"),
    ("preferred_change_method", "선호 변화관리 방식 (쉼표로 구분)"),
    ("preferred_communication", "선호 소통 방식 (쉼표로 구분)"),
    ("preferred_learning", "선호 학습 방식 (쉼표로 구분)"),
])
tip(doc, "앱 내 '입력 양식 샘플 다운로드'에서 1차 설문 CSV 샘플을 받을 수 있습니다.")

# 4-2
h(doc, "4-2. 2차 조직 자료 분석 (Excel 파일)", 2)
doc.add_paragraph("조직 환경을 분석하는 내부 자료입니다. Excel 파일에 6개 시트가 필요합니다.")
doc.add_paragraph()
table_2col(doc, [
    ("시트명", "분석 영역"),
    ("조직구조", "조직도, R&R, 의사결정 체계, 변화 추진 구조"),
    ("업무방식", "업무 프로세스, 회의 운영, 정보 공유, 문서 관리"),
    ("협업문화", "협업 구조, 부서 간 협업 수준, 인력 의존도"),
    ("제도정책", "KPI, 평가·보상 체계, 변화 활동 반영 여부"),
    ("디지털환경", "보유 솔루션, 라이선스 현황, 시스템 운영 정책"),
    ("변화경험", "과거 변화 프로젝트, 교육 이력, 변화 수용 경험"),
])
doc.add_paragraph()
doc.add_paragraph("각 시트의 필수 컬럼:")
table_2col(doc, [
    ("컬럼명", "설명"),
    ("평가항목", "진단 항목 이름"),
    ("평가점수", "1~5점 점수"),
    ("비고", "추가 메모 (선택)"),
])
tip(doc, "앱 내 샘플 다운로드에서 2차 조직자료 Excel 샘플을 받을 수 있습니다.")

# 4-3
h(doc, "4-3. 3차 인터뷰 진단 (JSON 파일)", 2)
doc.add_paragraph("임원·관리자·실무자·핵심이해관계자 인터뷰 결과 파일입니다.")
doc.add_paragraph()
table_2col(doc, [
    ("필드명", "설명"),
    ("interviewee_id", "인터뷰 대상자 ID"),
    ("role", "직급 (임원/관리자/실무자/핵심이해관계자)"),
    ("department", "부서명"),
    ("date", "인터뷰 날짜 (YYYY-MM-DD)"),
    ("pain_points", "현재 Pain Point"),
    ("resistance_factors", "변화 저항 원인 (배열)"),
    ("implicit_culture", "조직 내 암묵적 문화"),
    ("leadership_issues", "리더십 이슈"),
    ("success_experiences", "성공 경험"),
    ("failure_experiences", "실패 경험"),
    ("resistance_type", "저항 유형 (인지적/감정적/행동적)"),
    ("change_readiness_score", "변화 준비도 점수 (1~5점)"),
])
tip(doc, "앱 내 샘플 다운로드에서 3차 인터뷰 JSON 샘플을 받을 수 있습니다.")

# 4-4
h(doc, "4-4. 4차 행동 진단 (Excel 파일)", 2)
doc.add_paragraph("실제 행동 데이터를 분석합니다. Excel 파일에 4개 시트가 필요합니다.")
doc.add_paragraph()
table_2col(doc, [
    ("시트명", "주요 컬럼 (0~100% 수치)"),
    ("협업행동", "팀명, 회의참여율, 채팅활성도, 파일공유율, 공동편집비율"),
    ("디지털활용", "팀명, 라이선스사용률, 채팅활용률, 회의활용률, 활성사용자비율"),
    ("변화참여", "팀명, 교육참여율, 캠페인참여율, 챌린지참여율"),
    ("협업네트워크", "부서A, 부서B, 협업빈도점수, 정보공유점수"),
])
tip(doc, "앱 내 샘플 다운로드에서 4차 행동진단 Excel 샘플을 받을 수 있습니다.")
doc.add_page_break()

# ── 5. 결과 확인 ───────────────────────────────────────────
h(doc, "5. 진단 결과 확인")
doc.add_paragraph("분석 완료 후 화면에 아래 내용이 자동으로 출력됩니다.")
doc.add_paragraph()
table_2col(doc, [
    ("결과 항목", "내용"),
    ("종합 변화준비도", "0~100점 + 상/중/하 등급"),
    ("차수별 점수", "1~4차 각각의 점수 및 등급"),
    ("권고사항", "종합 등급에 따른 자동 출력"),
    ("1차 탭", "현재상태·변화추진환경 점수 + 해석 + 선호방식 집계"),
    ("2차 탭", "6개 영역별 점수 + 해석 텍스트"),
    ("3차 탭", "인터뷰 종합 점수 + 계층별 점수 + 저항유형 분포"),
    ("4차 탭", "4개 행동 영역별 점수 + 해석 텍스트"),
])
doc.add_paragraph()
h(doc, "등급 기준", 2)
table_2col(doc, [
    ("점수", "등급", ),
    ("80~100점", "상  —  변화 준비 완료"),
    ("60~79점",  "중  —  부분적 준비, 보완 필요"),
    ("0~59점",   "하  —  사전 준비 필요"),
])

# ── 6. 리포트 ──────────────────────────────────────────────
h(doc, "6. Word 리포트 다운로드")
doc.add_paragraph("분석 결과 화면 하단의 'Word 리포트 다운로드' 버튼을 클릭하면 자동 생성된 .docx 파일을 받을 수 있습니다.")
doc.add_paragraph()
table_2col(doc, [
    ("리포트 구성", "내용"),
    ("표지", "프로젝트명, 진단일, 폴라리스클라우드"),
    ("종합 변화준비도", "점수 바 + 차수별 요약표 + 권고사항"),
    ("1차 설문 분석", "현재상태·변화추진환경 점수 + 해석 + 요구사항"),
    ("2차 조직자료 분석", "6개 영역별 점수 + 해석"),
    ("3차 인터뷰 분석", "종합·계층별 점수 + 저항유형 분포"),
    ("4차 행동 진단", "4개 영역별 점수 + 해석"),
])
doc.add_page_break()

# ── 7. 코드 수정 ───────────────────────────────────────────
h(doc, "7. 코드 수정 방법 (VS Code)")
doc.add_paragraph("해석 텍스트나 권고사항을 수정하려면 VS Code에서 직접 편집합니다.")
doc.add_paragraph()
steps7 = [
    "VS Code 실행",
    "File → Open Folder → C:\\Claude\\polaris-diagnostic 선택",
    "왼쪽 탐색기에서 core → scoring.py 클릭",
    "INTERPRETATIONS 딕셔너리에서 원하는 영역·등급의 텍스트 수정",
    "RECOMMENDATIONS 딕셔너리에서 권고사항 수정",
    "Ctrl + S 저장",
]
for i, s in enumerate(steps7, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

doc.add_paragraph()
tip(doc, "scoring.py의 INTERPRETATIONS와 RECOMMENDATIONS만 수정하면 해석 텍스트와 권고사항이 바뀝니다. 다른 파일은 건드리지 않아도 됩니다.")

# ── 8. GitHub 업데이트 ─────────────────────────────────────
h(doc, "8. GitHub 업데이트 방법")
doc.add_paragraph("코드를 수정한 후 GitHub에 올리면 Streamlit 앱이 자동으로 업데이트됩니다.")
doc.add_paragraph()
steps8 = [
    "VS Code 터미널 열기 (Terminal → New Terminal)",
    "git add .",
    'git commit -m "수정 내용 간단히 설명"',
    "git push origin master",
    "Streamlit 앱이 자동으로 재배포됨 (1~2분 소요)",
]
for i, s in enumerate(steps8, 1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(s)
    if s.startswith("git"):
        r.font.name = "Courier New"
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

tip(doc, "코드를 수정할 때마다 git add . → git commit → git push 3단계를 반복합니다.")
doc.add_page_break()

# ── 9. FAQ ─────────────────────────────────────────────────
h(doc, "9. 자주 묻는 질문 (FAQ)")
faqs = [
    ("비밀번호를 바꾸고 싶어요.",
     "Streamlit Cloud 앱 설정 → Secrets에서 APP_PASSWORD 값을 변경하세요."),
    ("샘플 파일은 어디서 받나요?",
     "앱 진단실행 화면의 '입력 양식 샘플 다운로드' 버튼에서 4가지 파일을 모두 받을 수 있습니다."),
    ("해석 텍스트를 우리 상황에 맞게 바꾸고 싶어요.",
     "VS Code에서 core/scoring.py 파일의 INTERPRETATIONS 딕셔너리를 수정하면 됩니다."),
    ("새로운 진단 항목(질문)을 추가하고 싶어요.",
     "CSV/Excel 파일에 컬럼을 추가하고, scoring.py의 해당 함수에 컬럼명을 추가하면 됩니다."),
    ("앱이 느리거나 접속이 안 돼요.",
     "Streamlit Cloud 무료 플랜은 비활성 시 슬립 모드로 전환됩니다. 첫 접속 시 30초~1분 기다리면 됩니다."),
    ("SharePoint에 자동 저장은 언제 되나요?",
     "현재는 Word 다운로드만 가능합니다. SharePoint 연동은 Azure AD 앱 등록 후 .env에 값을 입력하면 활성화됩니다."),
    ("코드를 수정했는데 앱에 반영이 안 돼요.",
     "git push까지 완료했는지 확인하세요. Streamlit Cloud는 GitHub에 push해야 반영됩니다."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(f"Q. {q}").font.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"A. {a}")
    r.font.color.rgb = GRAY
    doc.add_paragraph()

doc.save("C:\\Claude\\폴라리스클라우드_진단에이전트_매뉴얼.docx")
print("매뉴얼 생성 완료!")
