"""
진단 결과 Word 리포트 자동 생성
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x00, 0x4B, 0x87)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1A, 0x6B, 0x3C)
GRAY  = RGBColor(0x59, 0x59, 0x59)

LEVEL_COLOR = {"상": "D4EDDA", "중": "FFF3CD", "하": "F8D7DA"}
LEVEL_TEXT_COLOR = {"상": GREEN, "중": RGBColor(0x85, 0x64, 0x04), "하": RED}


def _bg(cell, hex_color: str):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pr.append(shd)


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.color.rgb = BLUE
    r.font.bold = True


def _score_row(doc, label: str, score: float, level: str):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(8)
    t.columns[2].width = Cm(2.5)

    c0 = t.cell(0, 0)
    c0.text = label
    c0.paragraphs[0].runs[0].font.bold = True
    _bg(c0, "004B87")
    c0.paragraphs[0].runs[0].font.color.rgb = WHITE

    c1 = t.cell(0, 1)
    filled = int(score / 100 * 20)
    c1.text = "█" * filled + "░" * (20 - filled) + f"  {score}점"
    _bg(c1, LEVEL_COLOR.get(level, "FFFFFF"))

    c2 = t.cell(0, 2)
    c2.text = level
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c2.paragraphs[0].runs[0]
    r.font.bold = True
    r.font.color.rgb = LEVEL_TEXT_COLOR.get(level, GRAY)
    _bg(c2, LEVEL_COLOR.get(level, "FFFFFF"))
    doc.add_paragraph()


def generate(project_name: str, survey: dict, internal: dict,
             interview: dict, behavior: dict, synthesis: dict) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────
    doc.add_paragraph()
    for text, size, bold in [
        ("변화관리 초기 진단 리포트", 26, True),
        (f"프로젝트: {project_name}", 16, False),
        (datetime.now().strftime("%Y년 %m월 %d일"), 12, False),
        ("폴라리스클라우드", 14, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = BLUE
    doc.add_page_break()

    # ── 종합 요약 ──────────────────────────────────────────
    _h(doc, "종합 변화준비도")
    overall = synthesis.get("종합점수", 0)
    level   = synthesis.get("종합등급", "")
    _score_row(doc, "종합 변화준비도", overall, level)

    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, (label, key) in enumerate([
        ("1차 설문", "1차 설문"),
        ("2차 조직자료", "2차 조직자료"),
        ("3차 인터뷰", "3차 인터뷰"),
        ("4차 행동진단", "4차 행동진단"),
    ]):
        c = t.cell(0, i)
        score = synthesis["차수별_점수"].get(key, 0)
        lv = "상" if score >= 80 else "중" if score >= 60 else "하"
        c.text = f"{label}\n{score}점 ({lv})"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bg(c, LEVEL_COLOR.get(lv, "FFFFFF"))
    doc.add_paragraph()

    # 권고사항
    _h(doc, "종합 권고사항", 2)
    for rec in synthesis.get("권고사항", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(rec)
        r.font.bold = True
    doc.add_paragraph()
    doc.add_page_break()

    # ── 1차 설문 ───────────────────────────────────────────
    _h(doc, "1차 사전 진단 설문")

    _score_row(doc, "현재 상태 진단", survey.get("현재상태_점수", 0), survey.get("현재상태_등급", ""))
    doc.add_paragraph(survey.get("현재상태_해석", ""))
    doc.add_paragraph()

    _score_row(doc, "변화 추진 환경 진단", survey.get("변화추진환경_점수", 0), survey.get("변화추진환경_등급", ""))
    doc.add_paragraph(survey.get("변화추진환경_해석", ""))
    doc.add_paragraph()

    # 요구사항
    _h(doc, "요구사항 진단 결과", 2)
    for label, key in [
        ("선호하는 변화관리 방식", "선호_변화방식"),
        ("선호하는 소통 방식", "선호_소통방식"),
        ("선호하는 학습 방식", "선호_학습방식"),
    ]:
        data = survey.get(key, {})
        if data:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").font.bold = True
            p.add_run("  |  ".join(f"{k}({v}명)" for k, v in data.items()))
    doc.add_paragraph()
    doc.add_page_break()

    # ── 2차 조직자료 ───────────────────────────────────────
    _h(doc, "2차 조직 자료 분석 (환경 진단)")
    for area in ["조직구조", "업무방식", "협업문화", "제도정책", "디지털환경", "변화경험"]:
        data = internal.get(area, {})
        if data:
            _score_row(doc, area, data.get("점수", 0), data.get("등급", ""))
            doc.add_paragraph(data.get("해석", ""))
            doc.add_paragraph()
    doc.add_page_break()

    # ── 3차 인터뷰 ─────────────────────────────────────────
    _h(doc, "3차 인터뷰 진단 (심층 진단)")
    _score_row(doc, "인터뷰 종합", interview.get("종합점수", 0), interview.get("종합등급", ""))
    doc.add_paragraph(interview.get("해석", ""))
    doc.add_paragraph()

    # 계층별 점수
    role_scores = interview.get("계층별_점수", {})
    if role_scores:
        _h(doc, "계층별 변화준비도", 2)
        t = doc.add_table(rows=1, cols=len(role_scores))
        t.style = "Table Grid"
        for i, (role, score) in enumerate(role_scores.items()):
            lv = "상" if score >= 80 else "중" if score >= 60 else "하"
            c = t.cell(0, i)
            c.text = f"{role}\n{score}점 ({lv})"
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bg(c, LEVEL_COLOR.get(lv, "FFFFFF"))
        doc.add_paragraph()

    # 저항 유형
    resist = interview.get("저항유형_분포", {})
    if resist:
        p = doc.add_paragraph()
        p.add_run("저항 유형 분포: ").font.bold = True
        p.add_run("  |  ".join(f"{k}({v}건)" for k, v in resist.items()))
    doc.add_paragraph()
    doc.add_page_break()

    # ── 4차 행동 진단 ──────────────────────────────────────
    _h(doc, "4차 행동 진단")
    for area in ["협업행동", "디지털활용", "변화참여", "협업네트워크"]:
        data = behavior.get(area, {})
        if data:
            _score_row(doc, area, data.get("점수", 0), data.get("등급", ""))
            doc.add_paragraph(data.get("해석", ""))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
